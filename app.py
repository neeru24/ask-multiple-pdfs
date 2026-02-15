import streamlit as st
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document
import base64
import streamlit.components.v1 as components

from langchain_text_splitters import CharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaLLM


# -------------------------------
# MULTI-FORMAT TEXT EXTRACTION
# -------------------------------
def get_all_text(uploaded_files):
    text = ""

    for file in uploaded_files:
        file_extension = file.name.split(".")[-1].lower()

        if file_extension == "pdf":
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"

        elif file_extension == "docx":
            doc = Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"

        elif file_extension in ["txt", "md"]:
            text += file.read().decode("utf-8") + "\n"

        else:
            st.warning(f"Unsupported file type: {file.name}")

    return text


# -------------------------------
# TEXT SPLITTING
# -------------------------------
def get_text_chunks(text):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=200,
        chunk_overlap=20,
        length_function=len
    )
    return text_splitter.split_text(text)


# -------------------------------
# VECTOR STORE
# -------------------------------
def get_vectorstore(text_chunks):
    embeddings = OllamaEmbeddings(
        model="nomic-embed-text"
    )

    vectorstore = FAISS.from_texts(
        texts=text_chunks,
        embedding=embeddings
    )

    return vectorstore



@st.cache_resource
def load_llm():
    return OllamaLLM(
        model="phi3:mini",
        temperature=0.2,
        num_predict=150, 
        keep_alive=-1
    )


 
#  LLm - general ai (Chat-gpt mode)

def ask_general_llm(question, language):
    llm = load_llm()

    if language == "Hindi":
        lang_instruction = "Answer in Hindi only."
    else:
        lang_instruction = "Answer in English only."

    prompt = f"""
You are an intelligent AI assistant.
{lang_instruction}

Answer clearly and completely.
Give structured responses when needed.

Question:
{question}
"""

    return llm.invoke(prompt)



# -------------------------------
#  ASK LLM (Rag mode if the working with documents)
# -------------------------------
def ask_llm(vectorstore, question, language):

    docs = vectorstore.similarity_search(question, k=2)
    context = "\n\n".join([doc.page_content for doc in docs]) if docs else "No context found."

    if language == "Hindi":
        lang_instruction = "Answer in Hindi only."
        not_found_text = "दस्तावेज़ में नहीं मिला।"
    else:
        lang_instruction = "Answer in English only."
        not_found_text = "Not found in document."

    prompt = f"""
You are a helpful assistant.
Answer ONLY using the context below.
If the answer is not found, say: "{not_found_text}"
{lang_instruction}

Context:
{context}

Question:
{question}

 Answer clearly and concisely:
"""

    llm = load_llm()
    return llm.invoke(prompt)




# -------------------------------
# MAIN APP
# -------------------------------
def main():
    load_dotenv()

    st.set_page_config(page_title="Chat with Documents", page_icon="📚")

    # ---------------- LANGUAGE TOGGLE ----------------
    if "language" not in st.session_state:
        st.session_state.language = "English"

    with st.sidebar:
        st.subheader("🌐 Language / भाषा")
        language = st.radio("Select Language", ["English", "Hindi"])
        st.session_state.language = language

    lang = st.session_state.language

    if lang == "Hindi":
        st.title("दस्तावेज़ों से चैट करें 📚")
    else:
        st.title("Chat with Multiple Documents 📚")

    # ---------------- SESSION STATE ----------------
    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None

    if "messages" not in st.session_state:
        st.session_state.messages = []

    # ---------------- SIDEBAR ----------------
    with st.sidebar:
        if lang == "Hindi":
            st.subheader("दस्तावेज़ अपलोड करें")
        else:
            st.subheader("Upload Documents")

        uploaded_files = st.file_uploader(
            "Upload PDF, DOCX, TXT, MD files" if lang=="English" else "PDF, DOCX, TXT, MD फ़ाइल अपलोड करें",
            type=["pdf", "docx", "txt", "md"],
            accept_multiple_files=True
        )

        # -------- FULL DOCUMENT VIEWER --------
        if uploaded_files:
            st.subheader("📄 View Documents")

            for file in uploaded_files:
                with st.expander(f"📂 {file.name}", expanded=False):
                    ext = file.name.split(".")[-1].lower()

                    try:
                        # PDF preview + download (FIXED FOR CHROME)
                        if ext == "pdf":
                            file.seek(0)
                            file_bytes = file.read()
                            b64 = base64.b64encode(file_bytes).decode('utf-8')
                            
                            # Using <embed> with markdown is more reliable in Chrome than <iframe> components
                            pdf_display = f'<embed src="data:application/pdf;base64,{b64}" width="100%" height="600" type="application/pdf">'
                            st.markdown(pdf_display, unsafe_allow_html=True)
                            
                            st.download_button("Download PDF", file_bytes, file_name=file.name)

                        # DOCX viewer
                        elif ext == "docx":
                            file.seek(0)
                            doc = Document(file)
                            full_text = "\n".join([p.text for p in doc.paragraphs])
                            st.text_area("DOCX Content", full_text, height=400)
                            st.download_button("Download DOCX", file.getvalue(), file_name=file.name)

                        # TXT / MD viewer
                        elif ext in ["txt", "md"]:
                            file.seek(0)
                            text = file.read().decode("utf-8")
                            st.text_area("Text Content", text, height=400)
                            st.download_button("Download File", text, file_name=file.name)

                        else:
                            st.info("File type not supported for direct view.")

                    except Exception as e:
                        st.error(f"Error reading file: {e}")

        # -----------------------------------

        if st.button("Process Documents" if lang=="English" else "दस्तावेज़ प्रोसेस करें"):
            if not uploaded_files:
                st.warning("Please upload at least one document." if lang=="English" else "कृपया कम से कम एक दस्तावेज़ अपलोड करें।")
            else:
                with st.spinner("Reading & indexing documents..." if lang=="English" else "दस्तावेज़ पढ़े जा रहे हैं..."):
                    raw_text = get_all_text(uploaded_files)

                    if not raw_text.strip():
                        st.error("No readable text found." if lang=="English" else "कोई पढ़ने योग्य टेक्स्ट नहीं मिला।")
                    else:
                        text_chunks = get_text_chunks(raw_text)
                        st.session_state.vectorstore = get_vectorstore(text_chunks)
                        st.success("✅ Documents processed successfully!" if lang=="English" else "✅ दस्तावेज़ सफलतापूर्वक प्रोसेस हो गए!")

        if st.button("Clear Chat" if lang=="English" else "चैट साफ़ करें"):
            st.session_state.messages = []
            st.success("Chat history cleared!" if lang=="English" else "चैट हिस्ट्री साफ़ कर दी गई!")

    # ---------------- DISPLAY CHAT HISTORY ----------------
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # ---------------- CHAT INPUT ----------------
    user_question = st.chat_input(
        "Ask a question about your documents..." if lang=="English" else "दस्तावेज़ से सवाल पूछें..."
    )

    if user_question:

        # if st.session_state.vectorstore is None:
        #     st.warning("Please upload and process documents first." if lang=="English" else "पहले दस्तावेज़ अपलोड और प्रोसेस करें।")
        #     st.stop()

        st.session_state.messages.append(
            {"role": "user", "content": user_question}
        )

        with st.chat_message("user"):
            st.markdown(user_question)

        with st.chat_message("assistant"):
            with st.spinner("Thinking..." if lang=="English" else "सोच रहा हूँ..."):
                if st.session_state.vectorstore is not None:
                    answer = ask_llm(st.session_state.vectorstore, user_question,lang)
                else:
                    answer = ask_general_llm(user_question,lang)

                st.markdown(answer)

        # Save assistant message
        st.session_state.messages.append(
            {"role": "assistant", "content": answer}
        )




if __name__ == "__main__":
    main()
